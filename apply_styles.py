"""
apply_styles.py
Post-processes the pandoc-generated thesis_draft.docx and applies:
  - Heading 1  : 20 pt, Bold, Calibri, page-break-before
  - Heading 2  : 16 pt, Bold, Calibri
  - Heading 3  : 14 pt, Bold, Calibri
  - Normal / Body : 12 pt, Calibri, justified, 1.15 line-spacing
  - Table text : 11 pt
  - Code blocks : 10 pt, Courier New
  - TOC entries : matching sizes
  - Page margins : 2.54 cm all sides (standard A4)
  - Header/footer : page numbers bottom-centre
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import sys
import os

INPUT  = "thesis_draft.docx"
OUTPUT = "thesis.docx"

FONT_BODY    = "Calibri"
FONT_CODE    = "Courier New"
COLOR_H1     = RGBColor(31,  56,  150)   # dark navy
COLOR_H2     = RGBColor(46,  116, 181)   # Coventry blue
COLOR_H3     = RGBColor(46,  116, 181)
COLOR_BODY   = RGBColor(0,   0,   0)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_run_font(run, name, size_pt, bold=False, italic=False, color=None):
    run.font.name        = name
    run.font.size        = Pt(size_pt)
    run.font.bold        = bold
    run.font.italic      = italic
    if color:
        run.font.color.rgb = color
    # Also set the East-Asian and complex-script fonts
    rPr = run._r.get_or_add_rPr()
    for tag in (qn("w:rFonts"),):
        el = rPr.find(tag)
        if el is None:
            el = OxmlElement(tag)
            rPr.insert(0, el)
        el.set(qn("w:ascii"),     name)
        el.set(qn("w:hAnsi"),     name)
        el.set(qn("w:cs"),        name if name != FONT_CODE else FONT_CODE)


def set_para_format(para, font_name, size_pt, bold=False, italic=False,
                    color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=0, space_after=8, line_spacing=1.15,
                    keep_with_next=False, page_break_before=False):
    pf = para.paragraph_format
    pf.alignment          = align
    pf.space_before       = Pt(space_before)
    pf.space_after        = Pt(space_after)
    pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing       = line_spacing
    pf.keep_with_next     = keep_with_next
    pf.page_break_before  = page_break_before

    for run in para.runs:
        set_run_font(run, font_name, size_pt, bold, italic, color)


def style_exists(doc, name):
    return any(s.name == name for s in doc.styles)


def get_or_create_style(doc, name, base_name=None):
    if style_exists(doc, name):
        return doc.styles[name]
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base_name and style_exists(doc, base_name):
        style.base_style = doc.styles[base_name]
    return style


def apply_style_definition(doc, style_name, font_name, size_pt,
                             bold=False, color=None,
                             space_before=0, space_after=8,
                             line_spacing=1.15,
                             align=WD_ALIGN_PARAGRAPH.LEFT):
    """Update a built-in style's default formatting."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    font = style.font
    font.name  = font_name
    font.size  = Pt(size_pt)
    font.bold  = bold
    if color:
        font.color.rgb = color

    pf = style.paragraph_format
    pf.alignment         = align
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line_spacing


def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def add_page_numbers(doc):
    """Add page number to the footer of all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear existing footer paragraphs
        for para in footer.paragraphs:
            para.clear()
        if not footer.paragraphs:
            footer.add_paragraph()
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        # Insert PAGE field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def is_code_block(para):
    """Detect code-fence paragraphs by style name or indentation."""
    sname = para.style.name.lower()
    return "code" in sname or "verbatim" in sname or "source" in sname


def is_table_style(para):
    sname = para.style.name.lower()
    return "table" in sname


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print(f"Loading {INPUT}...")
    doc = Document(INPUT)

    # --- Page margins ---
    set_page_margins(doc)

    # --- Update built-in style definitions ---
    # Heading 1
    apply_style_definition(doc, "Heading 1",
        font_name=FONT_BODY, size_pt=20, bold=True,
        color=COLOR_H1, space_before=24, space_after=6,
        line_spacing=1.0)
    # Heading 2
    apply_style_definition(doc, "Heading 2",
        font_name=FONT_BODY, size_pt=16, bold=True,
        color=COLOR_H2, space_before=18, space_after=6,
        line_spacing=1.0)
    # Heading 3
    apply_style_definition(doc, "Heading 3",
        font_name=FONT_BODY, size_pt=14, bold=True,
        color=COLOR_H3, space_before=12, space_after=4,
        line_spacing=1.0)
    # Normal / Default
    apply_style_definition(doc, "Normal",
        font_name=FONT_BODY, size_pt=12,
        space_before=0, space_after=8,
        line_spacing=1.15,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    # Body Text
    for bname in ("Body Text", "Default Paragraph Style", "First Paragraph"):
        apply_style_definition(doc, bname,
            font_name=FONT_BODY, size_pt=12,
            space_before=0, space_after=8,
            line_spacing=1.15,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    print("Applying per-paragraph formatting...")
    for para in doc.paragraphs:
        sname = para.style.name

        if sname.startswith("Heading 1"):
            set_para_format(para, FONT_BODY, 20, bold=True,
                            color=COLOR_H1,
                            space_before=24, space_after=6,
                            line_spacing=1.0, page_break_before=False)

        elif sname.startswith("Heading 2"):
            set_para_format(para, FONT_BODY, 16, bold=True,
                            color=COLOR_H2,
                            space_before=18, space_after=6,
                            line_spacing=1.0, keep_with_next=True)

        elif sname.startswith("Heading 3"):
            set_para_format(para, FONT_BODY, 14, bold=True,
                            color=COLOR_H3,
                            space_before=12, space_after=4,
                            line_spacing=1.0, keep_with_next=True)

        elif is_code_block(para):
            set_para_format(para, FONT_CODE, 10,
                            space_before=4, space_after=4,
                            line_spacing=1.0,
                            align=WD_ALIGN_PARAGRAPH.LEFT)

        elif sname.startswith("TOC"):
            # TOC entries keep their indent; just fix font
            for run in para.runs:
                run.font.name = FONT_BODY
                run.font.size = Pt(12)

        else:
            # All other paragraphs → body style
            set_para_format(para, FONT_BODY, 12,
                            space_before=0, space_after=8,
                            line_spacing=1.15,
                            align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # --- Fix table cell text ---
    print("Formatting tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT_BODY
                        run.font.size = Pt(11)
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after  = Pt(2)

    # --- Page numbers ---
    add_page_numbers(doc)

    print(f"Saving {OUTPUT}...")
    doc.save(OUTPUT)
    size_kb = os.path.getsize(OUTPUT) // 1024
    print(f"Done. {OUTPUT} ({size_kb} KB)")


if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    main()
